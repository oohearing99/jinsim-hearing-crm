"""
Word 문서 리포트 생성 모듈
고객용 보청기 만족도 예측 리포트를 .docx 형식으로 생성 (1페이지 줄글 요약)
"""

from io import BytesIO
from datetime import datetime
from typing import Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go

# 로고 파일 경로
LOGO_PATH = Path(__file__).parent.parent.parent / "assets" / "logo.png"


def build_report_docx(
    user_input_dict: dict,
    features: dict,
    score: int,
    satisfaction_level: str,
    summary_text: str,
    recommendations: list[str],
    breakdown: dict,
    chart_fig: Optional[go.Figure] = None,
    audiogram_fig: Optional[go.Figure] = None
) -> BytesIO:
    """
    고객용 Word 리포트 생성 (한 페이지 줄글 요약)

    Args:
        user_input_dict: 사용자 입력 원본 데이터
        features: 전처리된 특징 딕셔너리
        score: 예측 만족도 점수
        satisfaction_level: 만족도 등급
        summary_text: 요약 텍스트
        recommendations: 추천 사항 리스트
        breakdown: 점수 breakdown 딕셔너리
        chart_fig: Plotly 차트 (선택적)
        audiogram_fig: 청력도 Plotly 차트 (선택적)

    Returns:
        BytesIO: Word 문서 바이트 스트림
    """
    doc = Document()

    # 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 로고 추가
    if LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Cm(5))
        doc.add_paragraph()  # 로고 아래 공백

    # 제목
    title = doc.add_heading("보청기 만족도 예측 컨설팅", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 41, 55)

    # 고객명과 날짜
    customer_name = user_input_dict.get('customer_name', '고객')

    greeting_para = doc.add_paragraph()
    greeting_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    greeting_run = greeting_para.add_run(f"{customer_name}님을 위한 맞춤 분석")
    greeting_run.font.size = Pt(14)
    greeting_run.font.bold = True
    greeting_run.font.color.rgb = RGBColor(59, 130, 246)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()  # 공백

    # 인사말
    intro_para = doc.add_paragraph()
    intro_text = intro_para.add_run(
        f"{customer_name}님 안녕하세요. 청력검사를 받으신 뒤 이렇게 서면으로 인사를 드립니다. "
        f"{customer_name}님의 현재 청력 상태와 보청기 착용 시 예상되는 만족도에 대해 자세히 알려드리고자 합니다. "
        f"또한 {customer_name}님뿐만 아니라 주변 가족 분들께도 정보를 드리기 위해 이와 같은 맞춤 컨설팅 서비스를 제공하고 있습니다."
    )
    intro_text.font.size = Pt(10)
    intro_para.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    # 맞춤 컨설팅 서비스 섹션
    service_heading = doc.add_paragraph()
    service_heading_run = service_heading.add_run("[1:1 맞춤 만족도 예측 서비스]")
    service_heading_run.font.size = Pt(12)
    service_heading_run.font.bold = True
    service_heading_run.font.color.rgb = RGBColor(31, 41, 55)

    service_para = doc.add_paragraph()
    service_text = service_para.add_run(
        "저희는 각 고객님의 청력 형태와 유형, 생활 환경, 선호도를 정확히 분석하여 "
        "가장 최적화된 보청기 착용 계획을 제안해드리는 맞춤 컨설팅 서비스를 진행하고 있습니다. "
        "최근 다양하게 출시되고 있는 보청기 가운데, 각 고객님의 불편사항과 희망사항을 고려하여 "
        "현재의 부족한 부분을 보충하고 불편한 부분을 최소화시킬 수 있는 최신 기능이 탑재된 "
        "최적의 보청기를 추천해드리고 있습니다."
    )
    service_text.font.size = Pt(10)
    service_para.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    # 청력 분석 결과
    analysis_heading = doc.add_paragraph()
    analysis_heading_run = analysis_heading.add_run(f"{customer_name}님의 청력 분석 결과")
    analysis_heading_run.font.size = Pt(12)
    analysis_heading_run.font.bold = True
    analysis_heading_run.font.color.rgb = RGBColor(31, 41, 55)

    # 청력 정보 텍스트
    loss_level_map = {
        'normal': '정상',
        'mild': '경도',
        'moderate': '중등도',
        'moderately_severe': '중고도',
        'severe': '고도',
        'profound': '심도'
    }

    loss_level = loss_level_map.get(features.get('loss_level', 'moderate'), '중등도')
    left_pta = user_input_dict['audiogram_left_pta']
    right_pta = user_input_dict['audiogram_right_pta']
    left_speech = user_input_dict['speech_score_left']
    right_speech = user_input_dict['speech_score_right']
    age = user_input_dict['age']

    analysis_para = doc.add_paragraph()
    analysis_text = analysis_para.add_run(
        f"{customer_name}님의 청력은 {datetime.now().strftime('%Y년 %m월')}에 평가한 결과, "
        f"우측 {right_pta:.0f}dB HL, 좌측 {left_pta:.0f}dB HL로 {loss_level} 난청을 보이고 있습니다. "
        f"어음명료도는 좌측 {left_speech}%, 우측 {right_speech}%로 측정되었습니다. "
        f"정상청력은 20dB 이내이며, 청력은 변화될 수 있으므로 반드시 정기적으로(연 1회) "
        f"청력평가를 받아 청력의 변화를 관리하셔야 합니다."
    )
    analysis_text.font.size = Pt(10)
    analysis_para.paragraph_format.line_spacing = 1.5

    # 청력도 이미지 삽입
    if audiogram_fig is not None:
        try:
            doc.add_paragraph()
            audiogram_heading = doc.add_paragraph()
            audiogram_heading_run = audiogram_heading.add_run("[청력도 (Audiogram)]")
            audiogram_heading_run.font.size = Pt(11)
            audiogram_heading_run.font.bold = True

            # Plotly 차트를 이미지로 변환
            audiogram_bytes = audiogram_fig.to_image(format="png", width=600, height=400, scale=2)
            audiogram_stream = BytesIO(audiogram_bytes)

            audiogram_para = doc.add_paragraph()
            audiogram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            audiogram_run = audiogram_para.add_run()
            audiogram_run.add_picture(audiogram_stream, width=Cm(14))
        except Exception as e:
            # 이미지 변환 실패 시 무시
            pass

    doc.add_paragraph()

    # 만족도 예측 결과
    prediction_para = doc.add_paragraph()
    prediction_text = prediction_para.add_run(
        f"청력은 한 번 떨어지기 시작하면 더 이상 좋아지지 않습니다. "
        f"그리하여 청력의 손실을 보완하는 방법으로 많은 사람들이 보청기를 사용하게 됩니다. "
        f"보청기는 떨어지는 청신경에 자극을 주어 청력이 떨어지는 것을 보완해주는 역할을 하게 됩니다. "
    )
    prediction_text.font.size = Pt(10)
    prediction_para.paragraph_format.line_spacing = 1.5

    # 만족도 점수 강조
    score_para = doc.add_paragraph()
    score_intro = score_para.add_run(
        f"\n현재 {customer_name}님의 청력 상태와 생활 환경, 선호도를 종합적으로 분석한 결과, "
        f"보청기 착용 시 예상 만족도는 "
    )
    score_intro.font.size = Pt(10)

    score_value = score_para.add_run(f"100점 만점에 {score}점")
    score_value.font.size = Pt(12)
    score_value.font.bold = True
    if score >= 85:
        score_value.font.color.rgb = RGBColor(16, 185, 129)
    elif score >= 70:
        score_value.font.color.rgb = RGBColor(59, 130, 246)
    elif score >= 55:
        score_value.font.color.rgb = RGBColor(245, 158, 11)
    else:
        score_value.font.color.rgb = RGBColor(239, 68, 68)

    score_outro = score_para.add_run(f"으로 '{satisfaction_level}' 수준입니다. ")
    score_outro.font.size = Pt(10)
    score_para.paragraph_format.line_spacing = 1.5

    # 요약 내용 (간결하게)
    summary_sentences = summary_text.split('. ')
    key_summary = '. '.join(summary_sentences[:3]) + '.'

    summary_para = doc.add_paragraph()
    summary_para.add_run(key_summary)
    for run in summary_para.runs:
        run.font.size = Pt(10)
    summary_para.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    # 착용 계획
    fitting_plan = user_input_dict.get('fitting_plan', 'bilateral')
    fitting_heading = doc.add_paragraph()
    fitting_heading_run = fitting_heading.add_run("[보청기 착용 계획]")
    fitting_heading_run.font.size = Pt(11)
    fitting_heading_run.font.bold = True

    fitting_para = doc.add_paragraph()
    if fitting_plan == 'bilateral':
        fitting_text = fitting_para.add_run(
            f"{customer_name}님께서는 양측 착용을 계획하고 계십니다. "
            "양측 착용은 양쪽 귀에 모두 보청기를 착용하는 방식으로, "
            "소리의 방향감 유지, 소음 환경에서 청취력 향상, 양쪽 귀의 균형적인 자극으로 청각 기능 유지, "
            "더 자연스러운 청취 경험 등의 장점이 있습니다. 양측 난청의 경우 양측 착용을 권장합니다."
        )
    else:
        side = "좌측" if fitting_plan == 'unilateral_left' else "우측"
        fitting_text = fitting_para.add_run(
            f"{customer_name}님께서는 {side} 단측 착용을 계획하고 계십니다. "
            "단측 착용은 초기 비용 부담이 적고 한쪽 귀에만 집중하여 적응할 수 있다는 장점이 있습니다. "
            "다만 소리의 방향감 저하, 소음 환경에서 청취력 감소, "
            "착용하지 않은 귀의 청각 기능 저하 가능성 등의 제한사항이 있을 수 있습니다. "
            "양측 난청이 있는 경우, 가능하다면 양측 착용으로 전환하는 것을 고려해보시길 권장합니다."
        )
    fitting_text.font.size = Pt(9)
    fitting_para.paragraph_format.line_spacing = 1.4

    doc.add_paragraph()

    # 전문가 추천사항
    rec_heading = doc.add_paragraph()
    rec_heading_run = rec_heading.add_run("[전문가 추천사항]")
    rec_heading_run.font.size = Pt(11)
    rec_heading_run.font.bold = True

    rec_para = doc.add_paragraph()
    rec_text_content = " ".join(recommendations[:3])  # 최대 3개를 하나의 문단으로
    rec_text = rec_para.add_run(rec_text_content)
    rec_text.font.size = Pt(9)
    rec_para.paragraph_format.line_spacing = 1.4

    doc.add_paragraph()

    # 시범 착용 안내
    trial_para = doc.add_paragraph()
    trial_text = trial_para.add_run(
        "보청기를 착용하실 경우, 1개월 시범 착용을 통해 직접 경험해보시고 "
        "마음에 드실 경우 착용하시면 되고, 만족스럽지 않을 경우 반품하실 수 있습니다. "
        "불편사항을 모두 고려하여 최신 기술이 탑재된 보청기를 제안 드리겠습니다."
    )
    trial_text.font.size = Pt(9)
    trial_text.font.bold = True
    trial_para.paragraph_format.line_spacing = 1.4

    # 추가 정보 (간략하게)
    main_complaints = user_input_dict.get('main_complaints', [])
    wearing_goal = user_input_dict.get('wearing_goal')

    if main_complaints or wearing_goal:
        doc.add_paragraph()
        extra_para = doc.add_paragraph()

        if main_complaints:
            complaints_text = extra_para.add_run(
                f"현재 주요 불편사항: {', '.join(main_complaints[:2])}. "
            )
            complaints_text.font.size = Pt(8)

        if wearing_goal:
            goal_text = extra_para.add_run(f"착용 목표: {wearing_goal}.")
            goal_text.font.size = Pt(8)

    doc.add_paragraph()

    # 면책 조항
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer.add_run(
        "※ 본 리포트는 참고 자료이며 의학적 진단이 아닙니다. "
        "최종 결정은 청각 전문가와 상담 후 내리시기 바랍니다."
    )
    disclaimer_run.font.size = Pt(7)
    disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)
    disclaimer_run.italic = True

    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
